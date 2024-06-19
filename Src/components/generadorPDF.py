import os
from tkinter import filedialog, Tk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx2pdf import convert
from PyPDF2 import PdfWriter, PdfReader
from reportlab.lib.pagesizes import legal
from reportlab.pdfgen import canvas
from io import BytesIO
from datetime import datetime
import locale

# Establecer la configuración regional a español
locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')

# Obtener el directorio actual y el directorio del proyecto
current_dir = os.path.dirname(os.path.realpath(__file__))
project_dir = os.path.dirname(current_dir)

# Ruta del archivo .docx original y del archivo modificado
input_data_path = os.path.join(
    project_dir, 'pdf', 'fga152_asistencia_clase.docx')
output_data_path = os.path.join(
    project_dir, 'pdf', 'fga152_asistencia_clase_modificado.docx')


def insert_data_into_table(doc_path, output_path, data, attended_docs):
    # Abrir el documento original
    doc = Document(doc_path)

    # Asumir que la tabla es la primera en el documento
    table = doc.tables[0]

    # Limpiar las filas existentes (excepto la fila de encabezado) y luego insertar datos
    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._element)

    # Longitud máxima permitida para el nombre
    max_length = 22

    # Insertar los datos en la tabla, comenzando desde la segunda fila
    for idx, (doc_num, name, program) in enumerate(data):
        # Cortar el último apellido si el nombre excede la longitud máxima
        if len(name) > max_length:
            name_parts = name.split()
            name_lines = []
            current_line = ''

            for part in name_parts:
                if len(current_line) + len(part) + 1 <= max_length:
                    if current_line:
                        current_line += ' '
                    current_line += part
                else:
                    name_lines.append(current_line)
                    current_line = part

            if current_line:
                name_lines.append(current_line)

            name = '\n'.join(name_lines)

        row_cells = table.add_row().cells

        # Establecer y formatear el número en la primera celda
        number_paragraph = row_cells[0].paragraphs[0]
        number_paragraph.clear()  # Limpiar el contenido existente
        # Alinear el párrafo al centro
        number_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        number_run = number_paragraph.add_run()
        number_run.bold = True
        number_run.font.size = Pt(11)
        number_run.font.name = 'Arial'
        number_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        number_run.text = str(idx + 1)

        # Insertar y formatear los datos en las demás celdas
        # Limpiar el contenido existente en la celda
        row_cells[1].paragraphs[0].clear()
        name_paragraph = row_cells[1].paragraphs[0]
        name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        name_run = name_paragraph.add_run()
        name_run.font.size = Pt(11)
        name_run.font.name = 'Arial'
        name_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        name_run.text = name

        row_cells[2].text = str(doc_num)
        row_cells[3].text = program

        # Formatear la celda de firma con centrado y texto correspondiente
        signature_paragraph = row_cells[4].paragraphs[0]
        signature_paragraph.clear()  # Limpiar el contenido existente
        # Alinear el párrafo al centro
        signature_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        signature_run = signature_paragraph.add_run()
        signature_run.font.size = Pt(11)
        signature_run.font.name = 'Arial'
        signature_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        signature_run.text = 'ASISTIÓ' if doc_num in attended_docs else 'NO ASISTIÓ'

        # Aplicar el formato Arial 11 sin negrita a cada celda de la fila, excepto la primera columna
        for cell in row_cells[1:4]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Guardar el documento modificado
    doc.save(output_path)


def convert_docx_to_pdf(input_path, output_path):
    # Convertir el documento .docx a .pdf usando docx2pdf
    convert(input_path, output_path)


def add_fields_to_pdf(input_pdf_path, custom_output_pdf_path, fecha, grupo, aula, docente, curso, tema, sede):
    # Cortar el nombre del docente si supera los 28 caracteres
    max_docente_length = 28
    if len(docente) > max_docente_length:
        # Dividir el nombre en partes (asumiendo que el nombre está en formato "Nombre Apellido1 Apellido2")
        name_parts = docente.split()

        # Construir el nombre truncado sin el último apellido si es necesario
        truncated_name = ''
        current_length = 0

        for part in name_parts:
            if current_length + len(part) <= max_docente_length:
                if truncated_name:
                    truncated_name += ' '
                truncated_name += part
                current_length += len(part) + 1
            else:
                break

        docente = truncated_name

    # Crear un buffer para el nuevo contenido del PDF
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=legal)  # Cambiado a tamaño oficio

    # Definir las coordenadas para cada campo
    fields = [
        (f"{fecha}", 140, 897),
        (f"{grupo}", 375, 897),
        (f"{aula}", 470, 897),
        (f"{docente}", 150, 880),
        (f"{curso}", 375, 880),
        (f"{tema}", 130, 860),
        (f"{sede}", 130, 840)
    ]

    # Agregar cada campo al lienzo en las coordenadas especificadas
    for text, x, y in fields:
        can.drawString(x, y, text)

    # Guardar el contenido del lienzo en el buffer
    can.save()

    # Mover el contenido del buffer al principio
    packet.seek(0)

    # Leer el contenido del PDF original
    existing_pdf = PdfReader(open(input_pdf_path, "rb"))

    # Crear un nuevo PDF que combine el contenido existente con el nuevo contenido
    output = PdfWriter()
    new_pdf = PdfReader(packet)

    # Agregar las páginas del PDF original al nuevo PDF
    for page_num in range(len(existing_pdf.pages)):
        page = existing_pdf.pages[page_num]
        if page_num == 0:
            # Solo añadir el nuevo contenido a la primera página
            page.merge_page(new_pdf.pages[0])
        output.add_page(page)

    # Escribir el PDF combinado en el archivo de salida personalizado
    with open(custom_output_pdf_path, "wb") as outputStream:
        output.write(outputStream)


def save_pdf(grupo, aula, docente, curso, tema, sede, estudiantes, asistencia):
    # Función para guardar el PDF generado en una ubicación personalizada
    root = Tk()
    root.withdraw()  # Ocultar la ventana principal de tkinter

    # Mostrar el cuadro de diálogo para guardar archivo
    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])

    # Verificar si se seleccionó un archivo
    if file_path:
        # Obtener la fecha actual en español
        fecha = datetime.now().strftime("%d de %B de %Y")

        # Insertar datos en la tabla del documento modificado (si es necesario)
        insert_data_into_table(
            output_data_path, output_data_path, estudiantes, asistencia)

        # Convertir el documento modificado a PDF
        convert_docx_to_pdf(output_data_path, file_path)

        # Añadir los campos adicionales al PDF
        add_fields_to_pdf(file_path, file_path, fecha, grupo,
                          aula, docente, curso, tema, sede)

        # print(f"Documento PDF guardado en: {file_path}")


if __name__ == "__main__":

    pass
